import openai
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import qn as oxml_qn

# ===================== 1. 基础配置 =====================
client = openai.OpenAI(
    base_url="http://192.168.100.85:1234/v1",  # 替换为实际的API地址
    api_key="111"     # 替换为实际的API Key
)

CONFIG_PATH = "format_contents_config.json"
OUTPUT_PATH = "format_json.docx"
DOC_TOPIC = "AI自动化办公项目分析报告2"

# ===================== 2. 读取JSON配置（包含format_rules） =====================
def load_format_config(json_path):
    with open(json_path, "r", encoding="utf-8") as f:
        config = json.load(f)
    
    # 对齐方式映射
    align_map = {
        "LEFT": WD_PARAGRAPH_ALIGNMENT.LEFT,
        "CENTER": WD_PARAGRAPH_ALIGNMENT.CENTER,
        "RIGHT": WD_PARAGRAPH_ALIGNMENT.RIGHT,
        "JUSTIFY": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    }
    
    # 转换标题和段落的对齐方式
    for elem_type in ["H1", "H2", "H3", "P"]:
        if elem_type in config:
            align_str = config[elem_type]["alignment"].upper()
            config[elem_type]["alignment"] = align_map.get(align_str, WD_PARAGRAPH_ALIGNMENT.LEFT)
    
    # 转换表格单元格对齐方式
    if "TABLE" in config:
        table_align_str = config["TABLE"]["cell_alignment"].upper()
        config["TABLE"]["cell_alignment"] = align_map.get(table_align_str, WD_PARAGRAPH_ALIGNMENT.CENTER)
    
    return config

# 加载配置（包含format_rules）
FORMAT_CONFIG = load_format_config(CONFIG_PATH)
TABLE_CONFIG = FORMAT_CONFIG.get("TABLE", {})
TOC_CONFIG = FORMAT_CONFIG.get("TOC", {})
PROMPT_CONFIG = FORMAT_CONFIG.get("PROMPT", {})
FORMAT_RULES = PROMPT_CONFIG.get("format_rules", [])  # 从JSON读取format_rules

# ===================== 3. 自定义样式（基于JSON配置） =====================
def setup_doc_styles(doc):
    # 设置标题样式（H1/H2/H3）
    for i in [1, 2, 3]:
        style_name = f'Heading {i}'
        elem_type = f'H{i}'
        if elem_type not in FORMAT_CONFIG:
            continue  # 严格依赖JSON配置，不存在则跳过
        
        font_conf = FORMAT_CONFIG[elem_type]
        style = doc.styles[style_name]
        
        # 字体设置
        style.font.name = font_conf["font_name"]
        style.font.size = Pt(font_conf["font_size"])
        style.font.bold = font_conf["bold"]
        style.paragraph_format.space_after = Pt(font_conf["space_after"])
        style.paragraph_format.alignment = font_conf["alignment"]
        
        # 中文字体设置
        r_fonts = style._element.rPr.rFonts
        r_fonts.set(oxml_qn('w:eastAsia'), font_conf["font_name"])
        r_fonts.set(oxml_qn('w:ascii'), font_conf["font_name"])
        
        # 大纲级别设置（关键：修复层级嵌套问题）
        ppr = style._element.get_or_add_pPr()
        # 先移除已存在的outlineLvl（避免重复）
        for elem in ppr.findall('.//w:outlineLvl', namespaces=ppr.nsmap):
            ppr.remove(elem)
        # 重新添加大纲级别（H1=0, H2=1, H3=2）
        outline_lvl = OxmlElement('w:outlineLvl')
        outline_lvl.set(oxml_qn('w:val'), str(i-1))
        ppr.append(outline_lvl)

    # 设置目录样式（TOC 1/2/3）
    toc_styles = [
        ('TOC 1', "toc1_font", "toc1_size", "toc1_bold", 0),
        ('TOC 2', "toc2_font", "toc2_size", "toc2_bold", 24),
        ('TOC 3', "toc3_font", "toc3_size", "toc3_bold", 48)
    ]
    for name, font_key, size_key, bold_key, indent in toc_styles:
        if name not in doc.styles and all(key in TOC_CONFIG for key in [font_key, size_key, bold_key]):
            toc_style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            toc_style.font.name = TOC_CONFIG[font_key]
            toc_style.font.size = Pt(TOC_CONFIG[size_key])
            toc_style.font.bold = TOC_CONFIG[bold_key]
            toc_style.paragraph_format.first_line_indent = Pt(indent)
            toc_style.paragraph_format.space_after = Pt(0)
            
            # 中文字体设置
            r_fonts = toc_style._element.rPr.rFonts
            r_fonts.set(oxml_qn('w:eastAsia'), TOC_CONFIG[font_key])
            r_fonts.set(oxml_qn('w:ascii'), TOC_CONFIG[font_key])

# ===================== 4. 目录插入（修复层级关联） =====================
def insert_toc_compatible(doc):
    # 添加目录标题
    if TOC_CONFIG:
        toc_title_para = doc.add_paragraph()
        toc_title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        toc_run = toc_title_para.add_run(TOC_CONFIG.get("title", "目录"))
        toc_run.font.name = TOC_CONFIG.get("title_font", "微软雅黑")
        toc_run.font.size = Pt(TOC_CONFIG.get("title_size", 22))
        toc_run.bold = TOC_CONFIG.get("title_bold", True)
        
        # 中文字体设置
        r_fonts = toc_run.element.rPr.rFonts
        r_fonts.set(oxml_qn('w:eastAsia'), TOC_CONFIG.get("title_font", "微软雅黑"))
        r_fonts.set(oxml_qn('w:ascii'), TOC_CONFIG.get("title_font", "微软雅黑"))
    
    # 插入空行分隔
    doc.add_paragraph()
    
    # 修复TOC域参数（确保层级独立）
    toc_xml = parse_xml(f'''
    <w:p {nsdecls("w")}>
      <w:r>
        <w:fldChar w:fldCharType="begin" {nsdecls("w")}/>
      </w:r>
      <w:r>
        <w:instrText xml:space="preserve">TOC \\o "1-3" \\h \\z \\u \\n \\t "TOC 1,1,TOC 2,2,TOC 3,3"</w:instrText>
      </w:r>
      <w:r>
        <w:fldChar w:fldCharType="separate" {nsdecls("w")}/>
      </w:r>
      <w:r>
        <w:fldChar w:fldCharType="end" {nsdecls("w")}/>
      </w:r>
    </w:p>
    ''')
    
    # 关键：使用add_paragraph()创建占位符，避免层级嵌套
    toc_para = doc.add_paragraph()
    toc_para._element.getparent().replace(toc_para._element, toc_xml)
    
    # 目录后分页
    doc.add_page_break()

# ===================== 5. 生成Prompt（从JSON读取format_rules） =====================
def generate_prompt(topic):
    req = PROMPT_CONFIG.get("requirements", {})
    
    # 从JSON读取format_rules（若为空则用默认）
    format_rules = FORMAT_RULES if FORMAT_RULES else [
        "⚠️ 输出格式为核心要求，必须100%遵守：",
        "1. 输出内容：仅包含一个JSON数组，无任何其他文字、符号、注释、换行、markdown标记",
        "2. JSON数组元素格式：",
        "   - 一级标题：{\"type\":\"H1\",\"content\":\"标题内容（如：1. 执行摘要）\"}",
        "   - 二级标题：{\"type\":\"H2\",\"content\":\"标题内容（如：3. 各维度深度分析）\"}",
        "   - 三级标题：{\"type\":\"H3\",\"content\":\"标题内容（如：3.1 安防）\"}",
        "   - 正文段落：{\"type\":\"正文\",\"content\":\"正文内容（完整、连贯的文本）\"}",
        "   - 表格内容：{\"type\":\"TABLE\",\"content\":\"表格行1|表格行2|...\\n---|---|---\\n内容1|内容2|...\"}",
        "3. JSON示例（可直接被Python json.loads解析）：",
        "   [",
        "     {\"type\":\"H1\",\"content\":\"1. 执行摘要\"},",
        "     {\"type\":\"正文\",\"content\":\"AI自动化办公可显著提升企业文档处理效率，降低人力成本。\"},",
        "     {\"type\":\"TABLE\",\"content\":\"功能模块|工具选型|适用场景\\n---|---|---\\nWord生成|python-docx|办公自动化\"},",
        "     {\"type\":\"H1\",\"content\":\"3. 各维度深度分析\"},",  # 关键：H1级别，独立层级
        "     {\"type\":\"H2\",\"content\":\"3.1 安防\"}",
        "   ]"
    ]
    
    # 标题层级规则（从JSON读取）
    title_levels = PROMPT_CONFIG.get("title_levels", [])
    if title_levels:
        format_rules.append("\n📋 标题层级规则：")
        for level_conf in title_levels:
            level = level_conf["level"]
            level_name = level_conf["name"]
            examples = level_conf["examples"]
            # 强调H1级别独立，避免嵌套
            format_rules.append(f"   - {level_name}（{level}）：{examples}（{level}为独立层级，不嵌套在其他标题内）")
    
    # 表格格式规则（从JSON读取）
    table_conf = PROMPT_CONFIG.get("table_rule", {})
    if table_conf:
        format_rules.append("\n📊 表格格式规则：")
        format_rules.append(f"   - {table_conf.get('note', '列用|分隔，表头下必须加---分隔行')}")
        for table_line in table_conf.get("examples", []):
            format_rules.append(f"     {table_line}")
    
    # 数量要求（从JSON读取）
    requirements = [
        "\n🔍 数量强制要求：",
        f"   - H1至少{req.get('h1_min_count', 7)}个（包含所有示例标题，每个H1为独立层级）",
        f"   - H2至少{req.get('h2_min_count', 1)}个（必须包含「{req.get('h2_mandatory_title', '3. 各维度深度分析')}」，且为H1级别）",
        f"   - H3在「{req.get('h2_mandatory_title', '3. 各维度深度分析')}」下至少{req.get('h3_min_count_under_h2', 8)}个（包含所有示例标题）",
        f"   - 表格至少{req.get('table_min_count', 1)}个",
        f"   - 正文内容完整，符合分析报告逻辑"
    ]
    
    # 关键：强调层级独立性
    level_warning = [
        "\n❗ 重要层级要求：",
        "   - 「3. 各维度深度分析」必须是H1级别（一级标题），独立存在，不嵌套在「1. 执行摘要」内",
        "   - 所有H1标题都是顶级层级，相互独立",
        "   - H2/H3仅嵌套在所属H1标题下，不跨层级嵌套"
    ]
    
    # 最终Prompt
    prompt = f"""
    请生成一篇关于「{topic}」的分析报告，严格遵守以下所有规则（优先级：格式 > 层级 > 数量 > 内容）：
    {chr(10).join(format_rules + requirements + level_warning)}
    
    ❗ 最终输出检查项：
    1. 仅输出JSON数组，无任何多余内容（包括但不限于："以下是生成的内容"、注释、空行、```等）
    2. JSON语法正确，无中文乱码，可直接被Python json.loads()解析
    3. 「3. 各维度深度分析」为独立H1级别标题，不嵌套在其他标题内
    4. 所有标题编号格式、数量要求均满足
    """
    return prompt

# ===================== 6. 模型调用（优化JSON解析） =====================
def get_model_content(topic):
    prompt = generate_prompt(topic)
    try:
        response = client.chat.completions.create(
            model="qwen/qwen3-vl-8b",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,  # 降低随机性，确保层级正确
            max_tokens=4000,
            timeout=60
        )
        # 获取并清理模型输出
        raw_content = response.choices[0].message.content.strip()
        
        # 清理常见的多余内容
        if raw_content.startswith("```json"):
            raw_content = raw_content.replace("```json", "").replace("```", "").strip()
        raw_content = raw_content.lstrip("：").lstrip("：").strip()
        
        # 解析JSON
        content_list = json.loads(raw_content)
        print("大模型生成结果:",content_list)
        print("大模型生成类型:",type(content_list))
        # 验证并修复层级（确保3. 各维度深度分析是H1）
        fixed_content = []
        for item in content_list:
            if isinstance(item, dict) and item.get("content") == "3. 各维度深度分析":
                item["type"] = "H1"  # 强制设为H1
            fixed_content.append(item)
        return fixed_content
    
    except json.JSONDecodeError as e:
        print(f"\n❌ JSON解析失败：{e}")
        print(f"📜 模型原始输出：\n{raw_content}")
        try:
            clean_content = raw_content.replace("\n", "").replace("    ", "").strip()
            content_list = json.loads(clean_content)
            print("✅ 二次清理后解析成功")
            return content_list
        except:
            print("❌ 二次清理仍解析失败，返回空列表")
            return []
    except Exception as e:
        print(f"\n❌ 模型调用失败：{e}")
        return []

# ===================== 7. 表格渲染 =====================
def render_table(doc, table_text):
    if not TABLE_CONFIG:
        return
    
    table_lines = [line.strip() for line in table_text.split("\n") if line.strip()]
    if not table_lines:
        return
    
    header_line = None
    content_lines = []
    for i, line in enumerate(table_lines):
        if "---" in line and i > 0:
            header_line = table_lines[i-1]
            content_lines = table_lines[i+1:]
            break
    if not header_line:
        header_line = table_lines[0]
        content_lines = table_lines[1:]
    
    header_cells = [cell.strip() for cell in header_line.split("|") if cell.strip()]
    col_count = len(header_cells)
    if col_count == 0:
        return
    
    table = doc.add_table(rows=1, cols=col_count)
    table.style = 'Table Grid'
    
    # 表头设置
    header_cells_obj = table.rows[0].cells
    for i, cell_text in enumerate(header_cells):
        cell = header_cells_obj[i]
        run = cell.paragraphs[0].add_run(cell_text)
        run.font.name = TABLE_CONFIG["cell_font_name"]
        run.font.size = Pt(TABLE_CONFIG["cell_font_size"])
        run.bold = TABLE_CONFIG["header_bold"]
        run.element.rPr.rFonts.set(qn('w:eastAsia'), TABLE_CONFIG["cell_font_name"])
        cell.paragraphs[0].alignment = TABLE_CONFIG["cell_alignment"]
        
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TABLE_CONFIG["header_bg_color"]}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)
        cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{TABLE_CONFIG["row_height"]*20}"/>'))
    
    # 内容行设置
    for line in content_lines:
        row_cells = [cell.strip() for cell in line.split("|") if cell.strip()]
        row_cells += [""] * (col_count - len(row_cells))
        row = table.add_row().cells
        for i, cell_text in enumerate(row_cells[:col_count]):
            cell = row[i]
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.name = TABLE_CONFIG["cell_font_name"]
            run.font.size = Pt(TABLE_CONFIG["cell_font_size"])
            run.bold = False
            run.element.rPr.rFonts.set(qn('w:eastAsia'), TABLE_CONFIG["cell_font_name"])
            cell.paragraphs[0].alignment = TABLE_CONFIG["cell_alignment"]
            cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{TABLE_CONFIG["row_height"]*20}"/>'))

# ===================== 8. 渲染完整文档（修复层级嵌套） =====================
def render_full_document(content_list, output_path):
    if not content_list or not isinstance(content_list, list):
        print("❌ 无有效JSON数组内容可渲染")
        return
    
    doc = Document()
    setup_doc_styles(doc)
    insert_toc_compatible(doc)
    
    # 遍历JSON数组渲染内容（确保每个标题独立）
    current_h1 = None
    for idx, item in enumerate(content_list):
        if not isinstance(item, dict) or "type" not in item or "content" not in item:
            print(f"⚠️ 跳过第{idx+1}个无效元素：{item}")
            continue
        
        item_type = item.get("type", "").strip()
        content = item.get("content", "").strip()
        
        if not item_type or not content:
            print(f"⚠️ 跳过第{idx+1}个空内容元素")
            continue
        
        # 渲染H1标题（独立顶级层级）
        if item_type == "H1":
            if "H1" not in FORMAT_CONFIG:
                print(f"⚠️ H1样式未配置，跳过：{content}")
                continue
            config = FORMAT_CONFIG["H1"]
            # 关键：使用add_heading时确保level=1，且不嵌套
            para = doc.add_heading(level=1)
            para.add_run(content)  # 重新添加run，避免样式继承
            
            # 重置当前H1标记
            current_h1 = content
            
            # 应用H1样式
            run = para.runs[0]
            run.font.name = config["font_name"]
            run.font.size = Pt(config["font_size"])
            run.bold = config["bold"]
            r_fonts = run.element.rPr.rFonts
            r_fonts.set(oxml_qn('w:eastAsia'), config["font_name"])
            r_fonts.set(oxml_qn('w:ascii'), config["font_name"])
            
            para.alignment = config["alignment"]
            para.paragraph_format.space_after = Pt(config["space_after"])
        
        # 渲染H2标题（仅嵌套在当前H1下）
        elif item_type == "H2":
            if "H2" not in FORMAT_CONFIG:
                print(f"⚠️ H2样式未配置，跳过：{content}")
                continue
            config = FORMAT_CONFIG["H2"]
            para = doc.add_heading(level=2)
            para.add_run(content)
            
            run = para.runs[0]
            run.font.name = config["font_name"]
            run.font.size = Pt(config["font_size"])
            run.bold = config["bold"]
            r_fonts = run.element.rPr.rFonts
            r_fonts.set(oxml_qn('w:eastAsia'), config["font_name"])
            r_fonts.set(oxml_qn('w:ascii'), config["font_name"])
            
            para.alignment = config["alignment"]
            para.paragraph_format.space_after = Pt(config["space_after"])
        
        # 渲染H3标题
        elif item_type == "H3":
            if "H3" not in FORMAT_CONFIG:
                print(f"⚠️ H3样式未配置，跳过：{content}")
                continue
            config = FORMAT_CONFIG["H3"]
            para = doc.add_heading(level=3)
            para.add_run(content)
            
            run = para.runs[0]
            run.font.name = config["font_name"]
            run.font.size = Pt(config["font_size"])
            run.bold = config["bold"]
            r_fonts = run.element.rPr.rFonts
            r_fonts.set(oxml_qn('w:eastAsia'), config["font_name"])
            r_fonts.set(oxml_qn('w:ascii'), config["font_name"])
            
            para.alignment = config["alignment"]
            para.paragraph_format.space_after = Pt(config["space_after"])
        
        # 渲染正文
        elif item_type == "正文":
            if "P" not in FORMAT_CONFIG:
                print(f"⚠️ 正文样式未配置，跳过：{content[:20]}...")
                continue
            config = FORMAT_CONFIG["P"]
            para = doc.add_paragraph()
            run = para.add_run(content)
            
            run.font.name = config["font_name"]
            run.font.size = Pt(config["font_size"])
            run.bold = config["bold"]
            r_fonts = run.element.rPr.rFonts
            r_fonts.set(oxml_qn('w:eastAsia'), config["font_name"])
            r_fonts.set(oxml_qn('w:ascii'), config["font_name"])
            
            para.alignment = config["alignment"]
            para.paragraph_format.space_after = Pt(config["space_after"])
            para.paragraph_format.first_line_indent = Pt(config["first_line_indent"])
            para.paragraph_format.line_spacing = config["line_spacing"]
        
        # 渲染表格
        elif item_type == "TABLE":
            if not TABLE_CONFIG:
                print("⚠️ 表格样式未配置，跳过表格内容")
                continue
            render_table(doc, content)
        
        else:
            print(f"⚠️ 未知元素类型{item_type}，跳过：{content[:20]}...")
    
    # 保存文档
    doc.save(output_path)
    print(f"\n✅ 文档生成完成！路径：{output_path}")
    print("📌 目录更新步骤：")
    print("   1. 打开生成的Word文档")
    print("   2. 右键点击目录区域 → 选择「更新域」")
    print("   3. 选择「更新整个目录」→ 点击「确定」")

# ===================== 9. 主函数 =====================
if __name__ == "__main__":
    print("===== 基于JSON配置生成文档 =====")
    model_content = get_model_content(DOC_TOPIC)
    
    if not model_content:
        print("❌ 模型内容为空或解析失败，终止")
        exit(1)
    # render_full_document(model_content, OUTPUT_PATH)
    print("===== 生成完成 =====")