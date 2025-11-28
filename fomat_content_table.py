import openai
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.style import WD_STYLE_TYPE

# ===================== 1. 基础配置 =====================
client = openai.OpenAI(
    base_url="http://192.168.100.85:1234/v1",
    api_key="111"
)

CONFIG_PATH = "format_contents_table.json"
OUTPUT_PATH = "无模板_带目录表格文档更新.docx"
DOC_TOPIC = "AI自动化办公项目分析报告2"

# ===================== 2. 读取JSON配置 =====================
def load_format_config(json_path):
    with open(json_path, "r", encoding="utf-8") as f:
        config = json.load(f)
    
    align_map = {
        "LEFT": WD_PARAGRAPH_ALIGNMENT.LEFT,
        "CENTER": WD_PARAGRAPH_ALIGNMENT.CENTER,
        "RIGHT": WD_PARAGRAPH_ALIGNMENT.RIGHT,
        "JUSTIFY": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    }
    
    for elem_type in ["H1", "H2", "H3", "P"]:
        config[elem_type]["alignment"] = align_map.get(config[elem_type]["alignment"].upper(), WD_PARAGRAPH_ALIGNMENT.LEFT)
    config["TABLE"]["cell_alignment"] = align_map.get(config["TABLE"]["cell_alignment"].upper(), WD_PARAGRAPH_ALIGNMENT.CENTER)
    
    return config

FORMAT_CONFIG = load_format_config(CONFIG_PATH)
TABLE_CONFIG = FORMAT_CONFIG["TABLE"]
TOC_CONFIG = FORMAT_CONFIG["TOC"]

# ===================== 3. 自定义样式（兼容旧版本） =====================
def setup_doc_styles(doc):
    # 定义标题样式
    for i, style_name in enumerate(['Heading 1', 'Heading 2', 'Heading 3']):
        style = doc.styles[style_name]
        font_conf = FORMAT_CONFIG[f'H{i+1}']
        style.font.name = font_conf["font_name"]
        style.font.size = Pt(font_conf["font_size"])
        style.font.bold = font_conf["bold"]
        style.paragraph_format.space_after = Pt(font_conf["space_after"])
        style.paragraph_format.alignment = font_conf["alignment"]
        # 解决中文字体
        style._element.rPr.rFonts.set(qn('w:eastAsia'), font_conf["font_name"])

    # 定义目录样式（TOC 1/2/3）
    toc_styles = [
        ('TOC 1', TOC_CONFIG["toc1_font"], TOC_CONFIG["toc1_size"], TOC_CONFIG["toc1_bold"], 0),
        ('TOC 2', TOC_CONFIG["toc2_font"], TOC_CONFIG["toc2_size"], TOC_CONFIG["toc2_bold"], 24),
        ('TOC 3', TOC_CONFIG["toc3_font"], TOC_CONFIG["toc3_size"], TOC_CONFIG["toc3_bold"], 48)
    ]
    for name, font, size, bold, indent in toc_styles:
        if name not in doc.styles:
            toc_style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            toc_style.font.name = font
            toc_style.font.size = Pt(size)
            toc_style.font.bold = bold
            toc_style.paragraph_format.first_line_indent = Pt(indent)
            toc_style.paragraph_format.space_after = Pt(0)
            toc_style._element.rPr.rFonts.set(qn('w:eastAsia'), font)

# ===================== 4. 核心：手动插入目录（兼容旧版本） =====================
def insert_toc_compatible(doc):
    """手动插入目录XML，替代add_table_of_contents()"""
    # 目录标题
    toc_title = doc.add_heading(TOC_CONFIG["title"], level=1)
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    toc_run = toc_title.runs[0]
    toc_run.font.name = TOC_CONFIG["title_font"]
    toc_run.font.size = Pt(TOC_CONFIG["title_size"])
    toc_run.bold = TOC_CONFIG["title_bold"]
    toc_run.element.rPr.rFonts.set(qn('w:eastAsia'), TOC_CONFIG["title_font"])
    
    # 插入空行
    doc.add_paragraph()
    
    # 手动构建目录XML（支持3级目录）
    toc_xml = parse_xml(f'''
    <w:p {nsdecls("w")}>
      <w:r>
        <w:fldChar w:fldCharType="begin"/>
      </w:r>
      <w:r>
        <w:instrText xml:space="preserve">TOC \\o "1-3" \\h \\z \\u</w:instrText>
      </w:r>
      <w:r>
        <w:fldChar w:fldCharType="separate"/>
      </w:r>
      <w:r>
        <w:fldChar w:fldCharType="end"/>
      </w:r>
    </w:p>
    ''')
    # 将目录XML添加到文档中
    doc._body._element.append(toc_xml)
    
    # 目录后分页
    doc.add_page_break()

# ===================== 5. 生成Prompt =====================
# def generate_prompt(topic):
#     """完全基于JSON配置循环拼接Prompt，无硬编码示例"""
#     # 从配置中读取核心参数
#     prompt_conf = FORMAT_CONFIG["PROMPT"]
#     req = prompt_conf["requirements"]
    
#     # ========== 1. 循环拼接标题层级规则 ==========
#     format_rules = [prompt_conf["title"], "1. 标题层级与编号严格匹配以下示例："]
#     # 遍历所有标题层级（H1/H2/H3）
#     for level_conf in prompt_conf["title_levels"]:
#         level = level_conf["level"]  # H1/H2/H3
#         level_name = level_conf["name"]  # 一级标题/二级标题
#         examples = level_conf["examples"]  # 该层级的所有示例
        
#         # 拼接该层级的标记格式+示例
#         tag = FORMAT_CONFIG[level]["tag"]
#         end_tag = tag.replace("【", "【/")
        
#         # 先拼接基础格式说明
#         format_rules.append(f"   - {level_name}（{level}）：标记为{tag}内容{end_tag}")
#         # 循环拼接该层级的所有示例
#         for idx, example in enumerate(examples, 1):
#             format_rules.append(f"     {idx}. {tag}{example}{end_tag}")
    
#     # ========== 2. 拼接正文规则 ==========
#     p_tag = FORMAT_CONFIG["P"]["tag"]
#     p_end_tag = p_tag.replace("【", "【/")
#     format_rules.append(f"2. 正文（P）：标记为{p_tag}内容{p_end_tag}，示例：")
#     format_rules.append(f"   {p_tag}{prompt_conf['paragraph_example']}{p_end_tag}")
    
#     # ========== 3. 循环拼接表格规则 ==========
#     table_conf = prompt_conf["table_rule"]
#     table_start = TABLE_CONFIG["tag_start"]
#     table_end = TABLE_CONFIG["tag_end"]
#     format_rules.append(f"3. 表格（TABLE）：标记为{table_start}表格内容{table_end}，{table_conf['note']}")
#     format_rules.append(f"   {table_start}")
#     # 循环拼接表格示例的每一行
#     for table_line in table_conf["examples"]:
#         format_rules.append(f"   {table_line}")
#     format_rules.append(f"   {table_end}")
    
#     # ========== 4. 拼接核心要求 ==========
#     requirements = [
#         f"1. 标题数量要求：",
#         f"   - H1至少{req['h1_min_count']}个（需包含配置中的所有H1示例标题）；",
#         f"   - H2至少{req['h2_min_count']}个（必须包含「{req['h2_mandatory_title']}」）；",
#         f"   - H3在「{req['h2_mandatory_title']}」下至少{req['h3_min_count_under_h2']}个（需包含配置中的所有H3示例标题）；",
#         f"2. 表格要求：至少包含{req['table_min_count']}个表格，格式严格匹配示例；",
#         f"3. 输出要求：{req['output_rule']}。"
#     ]
    
#     # ========== 5. 最终拼接Prompt ==========
#     prompt = f"""
#     请生成一篇关于「{topic}」的分析报告，严格遵循以下规则：
#     {chr(10).join(format_rules)}
    
#     核心要求：
#     {chr(10).join(requirements)}
#     """
#     return prompt
def generate_prompt(topic):
    format_rules = [
        f"### 格式规则 ###",
        f"1. 标题层级与编号严格匹配以下示例：",
        f"   - 一级标题（H1）：{FORMAT_CONFIG['H1']['tag']}1. 执行摘要{FORMAT_CONFIG['H1']['tag'].replace('【','【/')}",
        f"   - 二级标题（H2）：{FORMAT_CONFIG['H2']['tag']}3. 各维度深度分析{FORMAT_CONFIG['H2']['tag'].replace('【','【/')}",
        f"   - 三级标题（H3）：{FORMAT_CONFIG['H3']['tag']}1. 3.1 安防{FORMAT_CONFIG['H3']['tag'].replace('【','【/')}",
        f"2. 正文：{FORMAT_CONFIG['P']['tag']}正文内容{FORMAT_CONFIG['P']['tag'].replace('【','【/')}",
        f"3. 表格：{TABLE_CONFIG['tag_start']}表格内容{TABLE_CONFIG['tag_end']}，列用|分隔，表头下加---分隔"
    ]
    
    prompt = f"""
    请生成一篇关于「{topic}」的分析报告，严格遵循以下规则：
    {chr(10).join(format_rules)}
    
    核心要求：
    1. H1至少7个，H2至少1个（3. 各维度深度分析），H3在H2下至少8个；
    2. 标题编号严格匹配示例，至少包含1个表格；
    3. 输出仅保留标记文本，无多余内容。
    """
    return prompt
能不能把标题/内容都放在配置文件
# ===================== 6. 模型调用+表格渲染 =====================
def get_model_content(topic):
    prompt = generate_prompt(topic)
    try:
        response = client.chat.completions.create(
            model="qwen/qwen3-vl-8b",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=4000,
            timeout=60
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"模型调用失败：{e}")
        return ""

def render_table(doc, table_text):
    table_lines = [line.strip() for line in table_text.split("\n") if line.strip()]
    if not table_lines:
        return
    
    header_line = None
    content_lines = []
    for i, line in enumerate(table_lines):
        if "---" in line and i > 0:
            header_line = table_lines[i-1]
            content_lines = table_lines[i+1:]
            break
    if not header_line:
        header_line = table_lines[0]
        content_lines = table_lines[1:]
    
    header_cells = [cell.strip() for cell in header_line.split("|") if cell.strip()]
    col_count = len(header_cells)
    if col_count == 0:
        return
    
    table = doc.add_table(rows=1, cols=col_count)
    table.style = 'Table Grid'
    
    # 表头
    header_cells_obj = table.rows[0].cells
    for i, cell_text in enumerate(header_cells):
        cell = header_cells_obj[i]
        run = cell.paragraphs[0].add_run(cell_text)
        run.font.name = TABLE_CONFIG["cell_font_name"]
        run.font.size = Pt(TABLE_CONFIG["cell_font_size"])
        run.bold = TABLE_CONFIG["header_bold"]
        run.element.rPr.rFonts.set(qn('w:eastAsia'), TABLE_CONFIG["cell_font_name"])
        cell.paragraphs[0].alignment = TABLE_CONFIG["cell_alignment"]
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TABLE_CONFIG["header_bg_color"]}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)
        cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{TABLE_CONFIG["row_height"]*20}"/>'))
    
    # 内容行
    for line in content_lines:
        row_cells = [cell.strip() for cell in line.split("|") if cell.strip()]
        row_cells += [""] * (col_count - len(row_cells))
        row = table.add_row().cells
        for i, cell_text in enumerate(row_cells[:col_count]):
            cell = row[i]
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.name = TABLE_CONFIG["cell_font_name"]
            run.font.size = Pt(TABLE_CONFIG["cell_font_size"])
            run.bold = False
            run.element.rPr.rFonts.set(qn('w:eastAsia'), TABLE_CONFIG["cell_font_name"])
            cell.paragraphs[0].alignment = TABLE_CONFIG["cell_alignment"]
            cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{TABLE_CONFIG["row_height"]*20}"/>'))

# ===================== 7. 渲染完整文档 =====================
def render_full_document(model_content, output_path):
    doc = Document()
    setup_doc_styles(doc)
    
    # 插入目录（兼容旧版本）
    insert_toc_compatible(doc)
    
    # 渲染正文/表格
    lines = model_content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        
        # 处理表格
        if line.startswith(TABLE_CONFIG["tag_start"]):
            table_text = []
            i += 1
            while i < len(lines):
                sub_line = lines[i].strip()
                if sub_line.endswith(TABLE_CONFIG["tag_end"]):
                    break
                table_text.append(sub_line)
                i += 1
            render_table(doc, "\n".join(table_text))
            i += 1
            continue
        
        # 处理标题/正文
        elem_type = None
        elem_text = None
        for type_key in ["H1", "H2", "H3", "P"]:
            config = FORMAT_CONFIG[type_key]
            start_tag = config["tag"]
            end_tag = config["tag"].replace("【", "【/")
            if line.startswith(start_tag) and line.endswith(end_tag):
                elem_type = type_key
                elem_text = line.replace(start_tag, "").replace(end_tag, "")
                break
        
        if not elem_type:
            elem_type = "P"
            elem_text = line
        
        # 应用格式
        config = FORMAT_CONFIG[elem_type]
        if "H" in elem_type:
            para = doc.add_heading(elem_text, level=int(elem_type[-1]))
        else:
            para = doc.add_paragraph()
            para.add_run(elem_text)
        
        run = para.runs[0]
        run.font.name = config["font_name"]
        run.font.size = Pt(config["font_size"])
        run.bold = config["bold"]
        run.element.rPr.rFonts.set(qn('w:eastAsia'), config["font_name"])
        
        para.alignment = config["alignment"]
        para.paragraph_format.space_after = Pt(config["space_after"])
        
        if elem_type == "P":
            para.paragraph_format.first_line_indent = Pt(config["first_line_indent"])
            para.paragraph_format.line_spacing = config["line_spacing"]
        
        i += 1
    
    doc.save(output_path)
    print(f"\n✅ 文档生成完成！路径：{output_path}")
    print("📌 打开文档后右键目录 → 「更新域」→ 「更新整个目录」即可显示完整目录。")

# ===================== 8. 主函数 =====================
if __name__ == "__main__":
    print("===== 无模板+兼容旧版本生成文档 =====")
    model_content = get_model_content(DOC_TOPIC)
    if not model_content:
        print("❌ 模型内容为空，终止")
        exit(1)
    render_full_document(model_content, OUTPUT_PATH)
    print("===== 生成完成 =====")