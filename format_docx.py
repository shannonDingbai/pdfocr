### 使用大模型生成文本
### 配置文本格式
### 使用python-docx docx文件



import openai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# ===================== 新版 openai 配置（核心修改）=====================
# 1. 初始化客户端（替代旧版的直接配置 api_base/api_key）
client = openai.OpenAI(
    base_url="http://192.168.100.85:1234/v1",  # 本地模型地址（LM Studio/Ollama）
    api_key="123"  # 本地模型填任意非空值即可
)

# 格式配置（不变）
FORMAT_CONFIG = {
    "H1": {"tag": "【H1】", "font_name": "微软雅黑", "font_size": 22, "bold": True, "space_after": 12, "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER},
    "H2": {"tag": "【H2】", "font_name": "微软雅黑", "font_size": 15, "bold": True, "space_after": 8, "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT},
    "P": {"tag": "【P】", "font_name": "宋体", "font_size": 12, "bold": False, "first_line_indent": 24, "line_spacing": 1.5, "space_after": 0, "alignment": WD_PARAGRAPH_ALIGNMENT.JUSTIFY}
}

def generate_prompt(topic, format_config):
    """动态生成含格式配置的Prompt（不变）"""
    format_rules = []
    for elem_type, config in format_config.items():
        if elem_type == "P":
            format_rules.append(
                f"- {elem_type}（正文）：标记为{config['tag']}内容{config['tag'].replace('【', '【/')}，"
                f"字体{config['font_name']}，{config['font_size']}磅，首行缩进{config['first_line_indent']//12}字符，行间距{config['line_spacing']}倍"
            )
        else:
            format_rules.append(
                f"- {elem_type}（标题）：标记为{config['tag']}内容{config['tag'].replace('【', '【/')}，"
                f"字体{config['font_name']}，{config['font_size']}磅，{'加粗' if config['bold'] else '不加粗'}，段后间距{config['space_after']}磅"
            )
    
    prompt = f"""
    请生成一篇关于「{topic}」的技术文档，严格按照以下规则输出结构化标记文本：
    {chr(10).join(format_rules)}
    要求：至少包含2个H1标题、3个H2标题，正文逻辑完整；输出仅保留标记文本，无多余内容。
    """
    return prompt

def get_structured_content(topic, format_config):
    """调用本地模型API（核心修改：新版接口）"""
    prompt = generate_prompt(topic, format_config)
    
    # 新版调用方式：client.chat.completions.create（替代旧版 openai.ChatCompletion.create）
    response = client.chat.completions.create(
        model="qwen/qwen3-vl-8b",  # 替换为你的本地模型名称
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
        max_tokens=2000
    )
    # 新版提取内容的方式（不变）
    return response.choices[0].message.content.strip()

# Word渲染逻辑（完全不变）
def render_word(content, format_config, save_path):
    doc = Document()
    lines = content.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        elem_type = None
        elem_text = None
        for type_key, config in format_config.items():
            start_tag = config["tag"]
            end_tag = config["tag"].replace("【", "【/")
            if line.startswith(start_tag) and line.endswith(end_tag):
                elem_type = type_key
                elem_text = line.replace(start_tag, "").replace(end_tag, "")
                break
        
        if not elem_type:
            elem_type = "P"
            elem_text = line
        
        config = format_config[elem_type]
        if "H" in elem_type:
            para = doc.add_heading(level=int(elem_type[-1]))
        else:
            para = doc.add_paragraph()
        
        run = para.add_run(elem_text)
        run.font.name = config["font_name"]
        run.font.size = Pt(config["font_size"])
        run.bold = config["bold"]
        run.element.rPr.rFonts.set(qn('w:eastAsia'), config["font_name"])
        
        para.paragraph_format.space_after = Pt(config["space_after"])
        para.alignment = config["alignment"]
        if elem_type == "P":
            para.paragraph_format.first_line_indent = Pt(config["first_line_indent"])
            para.paragraph_format.line_spacing = config["line_spacing"]
    
    doc.save(save_path)
    print(f"文档已保存至：{save_path}")

# 执行生成
if __name__ == "__main__":
    structured_text = get_structured_content("AI自动化文档生成", FORMAT_CONFIG)
    print("本地模型返回内容：\n", structured_text)
    render_word(structured_text, FORMAT_CONFIG, "新版API生成文档.docx")